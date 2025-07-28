"""
Simplified OCR Resume Parser using PaddleOCR and spaCy
for structured field extraction from resumes.
"""

import logging
import re
import io
from typing import Optional, List, Dict
from dataclasses import dataclass
from pathlib import Path

# Core processing libraries
import numpy as np
from PIL import Image
import pdf2image
import PyPDF2
import docx

# OCR and NLP
try:
    from paddleocr import PaddleOCR
except ImportError:
    PaddleOCR = None

try:
    import pytesseract
except ImportError:
    pytesseract = None

try:
    import spacy
except ImportError:
    spacy = None

logger = logging.getLogger(__name__)

@dataclass
class WorkExperience:
    """Structured work experience entry"""
    company: str
    title: str
    location: Optional[str] = None
    start_date: Optional[str] = None
    end_date: Optional[str] = None
    description: Optional[str] = None
    duration: Optional[str] = None

@dataclass
class ResumeData:
    """Simplified structured resume data"""
    # Core Identity Fields
    name: str
    email: str
    phone: str
    
    # Professional Information
    title: str
    summary: str
    
    # Experience & Skills
    skills: List[str]
    experience: List[WorkExperience]
    education: str
    
    # Additional Fields
    location: str
    linkedin_url: Optional[str] = None
    github_url: Optional[str] = None
    years_experience: Optional[int] = None

class SimpleResumeParser:
    """Simplified OCR parser using PaddleOCR with spaCy NLP"""
    
    def __init__(self):
        self._initialize_ocr()
        self._initialize_nlp()
    
    def _initialize_ocr(self):
        """Initialize OCR engine with fallback"""
        self.ocr = None
        
        # Try PaddleOCR first
        if PaddleOCR:
            try:
                self.ocr = PaddleOCR(
                    use_angle_cls=True,
                    lang='en',
                    use_gpu=False,
                    show_log=False
                )
                logger.info("PaddleOCR initialized successfully")
            except Exception as e:
                logger.warning(f"PaddleOCR initialization failed: {e}")
        
        # Fallback to PyTesseract
        if not self.ocr and pytesseract:
            try:
                # Test pytesseract availability
                pytesseract.get_tesseract_version()
                self.use_tesseract = True
                logger.info("PyTesseract fallback ready")
            except Exception as e:
                logger.warning(f"PyTesseract not available: {e}")
                self.use_tesseract = False
        else:
            self.use_tesseract = False
    
    def _initialize_nlp(self):
        """Initialize spaCy NLP model"""
        self.nlp = None
        
        if spacy:
            try:
                # Try loading different models
                for model_name in ["en_core_web_sm", "en_core_web_md", "en_core_web_lg"]:
                    try:
                        self.nlp = spacy.load(model_name)
                        logger.info(f"spaCy model '{model_name}' loaded successfully")
                        break
                    except OSError:
                        continue
                
                if not self.nlp:
                    logger.warning("No spaCy models found. Install with: python -m spacy download en_core_web_sm")
            except Exception as e:
                logger.warning(f"spaCy initialization failed: {e}")
    
    def parse_resume_file(self, file_content: bytes, filename: str) -> ResumeData:
        """
        Parse resume from uploaded file
        
        Args:
            file_content: Binary content of the file
            filename: Original filename
            
        Returns:
            ResumeData with structured fields
        """
        logger.info(f"Parsing resume: {filename}")
        
        # Extract text based on file type
        if filename.lower().endswith('.pdf'):
            text = self._extract_text_from_pdf(file_content)
        elif filename.lower().endswith(('.docx', '.doc')):
            text = self._extract_text_from_docx(file_content)
        elif filename.lower().endswith(('.png', '.jpg', '.jpeg', '.tiff', '.bmp')):
            text = self._extract_text_from_image(file_content)
        else:
            raise ValueError(f"Unsupported file type: {filename}")
        
        # Parse structured data from text
        parsed_data = self._parse_resume_text(text)
        
        logger.info(f"Successfully parsed resume for: {parsed_data.name}")
        return parsed_data
    
    def _extract_text_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF"""
        text = ""
        
        # Method 1: Direct text extraction
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text.strip():
                    text += page_text + "\n"
            
            # If we got substantial text, use it
            if len(text.strip()) > 100:
                return text
        except Exception as e:
            logger.warning(f"PyPDF2 extraction failed: {e}")
        
        # Method 2: OCR fallback
        try:
            images = pdf2image.convert_from_bytes(file_content, dpi=200, fmt='RGB')
            
            ocr_text = ""
            for image in images:
                page_text = self._extract_text_with_ocr(image)
                ocr_text += page_text + "\n"
            
            return ocr_text if ocr_text.strip() else text
        except Exception as e:
            logger.error(f"PDF OCR extraction failed: {e}")
            return text
    
    def _extract_text_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(io.BytesIO(file_content))
            text = ""
            
            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise
    
    def _extract_text_from_image(self, file_content: bytes) -> str:
        """Extract text from image using OCR"""
        try:
            image = Image.open(io.BytesIO(file_content))
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            return self._extract_text_with_ocr(image)
        except Exception as e:
            logger.error(f"Image OCR failed: {e}")
            raise
    
    def _extract_text_with_ocr(self, image: Image.Image) -> str:
        """Extract text using available OCR engine"""
        # Try PaddleOCR first
        if self.ocr:
            try:
                img_array = np.array(image)
                result = self.ocr.ocr(img_array)
                
                if result and result[0]:
                    text_blocks = []
                    for line in result[0]:
                        if line:
                            bbox, (text, confidence) = line
                            if confidence > 0.5:  # Only high-confidence text
                                text_blocks.append(text)
                    
                    return ' '.join(text_blocks)
            except Exception as e:
                logger.warning(f"PaddleOCR failed: {e}")
        
        # Fallback to PyTesseract
        if self.use_tesseract:
            try:
                return pytesseract.image_to_string(image)
            except Exception as e:
                logger.error(f"PyTesseract failed: {e}")
        
        logger.error("No OCR engine available")
        return ""
    
    def _parse_resume_text(self, text: str) -> ResumeData:
        """Parse structured data from resume text"""
        # Clean text
        text = self._clean_text(text)
        
        # Extract all fields
        return ResumeData(
            name=self._extract_name(text),
            email=self._extract_email(text),
            phone=self._extract_phone(text),
            title=self._extract_title(text),
            summary=self._extract_summary(text),
            skills=self._extract_skills(text),
            experience=self._extract_experience(text),
            education=self._extract_education(text),
            location=self._extract_location(text),
            linkedin_url=self._extract_linkedin_url(text),
            github_url=self._extract_github_url(text),
            years_experience=self._extract_years_experience(text)
        )
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove problematic characters
        text = re.sub(r'[^\w\s@.,()\-+/:#]', ' ', text)
        return text.strip()
    
    def _extract_name(self, text: str) -> str:
        """Extract name using spaCy NER and patterns"""
        # Method 1: spaCy NER
        if self.nlp:
            try:
                doc = self.nlp(text[:800])  # Check first 800 chars
                for ent in doc.ents:
                    if ent.label_ == "PERSON" and len(ent.text.split()) <= 4:
                        name = ent.text.strip()
                        if self._is_valid_name(name):
                            return name
            except Exception as e:
                logger.debug(f"spaCy name extraction failed: {e}")
        
        # Method 2: Pattern-based extraction
        lines = text.split('\n')
        for line in lines[:10]:  # Check first 10 lines
            line = line.strip()
            if not line or len(line) > 60:
                continue
            
            # Skip lines with contact info indicators
            if any(indicator in line.lower() for indicator in [
                '@', 'phone', 'email', 'address', 'linkedin', 'github', 'www', 'http'
            ]):
                continue
            
            # Check if it looks like a name
            words = line.split()
            if 2 <= len(words) <= 4:
                if all(word[0].isupper() and word.isalpha() for word in words):
                    name = ' '.join(words)
                    if self._is_valid_name(name):
                        return name
        
        return "Unknown Name"
    
    def _is_valid_name(self, name: str) -> bool:
        """Validate if a string is likely a person's name"""
        if not name or len(name.strip()) < 3:
            return False
        
        words = name.strip().split()
        if len(words) < 2 or len(words) > 4:
            return False
        
        # Check if all words are valid name parts
        for word in words:
            if not word.isalpha() or len(word) < 2:
                return False
        
        # Exclude common non-name patterns
        name_lower = name.lower()
        exclude_words = [
            'university', 'college', 'institute', 'engineer', 'developer', 
            'manager', 'analyst', 'consultant', 'director', 'resume'
        ]
        
        return not any(word in name_lower for word in exclude_words)
    
    def _extract_email(self, text: str) -> str:
        """Extract email address"""
        pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        matches = re.findall(pattern, text)
        return matches[0] if matches else ""
    
    def _extract_phone(self, text: str) -> str:
        """Extract phone number"""
        patterns = [
            r'\+\d{1,3}[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
            r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
            r'\d{3}[-.\s]?\d{3}[-.\s]?\d{4}'
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text)
            if matches:
                return matches[0]
        return ""
    
    def _extract_title(self, text: str) -> str:
        """Extract job title"""
        lines = text.split('\n')
        
        # Look for title indicators
        title_patterns = [
            r'(?:title|position|role):\s*(.+)',
            r'\b(?:Senior|Lead|Principal|Junior)?\s*(?:Software|Full[- ]?Stack|Frontend|Backend|Web|Mobile)?\s*(?:Engineer|Developer|Architect|Manager|Director)\b',
            r'\b(?:Product|Project|Engineering|Technical)\s*(?:Manager|Director|Lead)\b'
        ]
        
        # Check for explicit title indicators
        for line in lines[:15]:
            for pattern in title_patterns:
                match = re.search(pattern, line, re.IGNORECASE)
                if match:
                    if 'title' in pattern:
                        return match.group(1).strip()
                    else:
                        return match.group(0).strip()
        
        return ""
    
    def _extract_summary(self, text: str) -> str:
        """Extract summary/objective section"""
        patterns = [
            r'(?:Professional\s+)?Summary[:\s]*\n?(.*?)(?=\n\s*[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\s*:|\n\s*\n|\Z)',
            r'(?:Career\s+)?Objective[:\s]*\n?(.*?)(?=\n\s*[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\s*:|\n\s*\n|\Z)',
            r'Profile[:\s]*\n?(.*?)(?=\n\s*[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\s*:|\n\s*\n|\Z)'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE | re.DOTALL)
            if match:
                summary = match.group(1).strip()
                summary = re.sub(r'\s+', ' ', summary)
                if 50 <= len(summary) <= 500:
                    return summary
        
        return ""
    
    def _extract_skills(self, text: str) -> List[str]:
        """Extract skills using patterns and predefined lists"""
        skills = set()
        
        # Method 1: Find skills sections
        skills_patterns = [
            r'(?:Technical\s+)?Skills[:\s]*\n?(.*?)(?=\n\s*\n|\n[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\s*:|\Z)',
            r'Technologies[:\s]*\n?(.*?)(?=\n\s*\n|\n[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\s*:|\Z)'
        ]
        
        for pattern in skills_patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE | re.DOTALL)
            if match:
                skills_text = match.group(1)
                # Split by common delimiters
                skill_items = re.split(r'[,•\n\|;·\-\s]{2,}', skills_text)
                for skill in skill_items:
                    skill = skill.strip()
                    if 2 <= len(skill) <= 25 and skill.replace(' ', '').replace('.', '').replace('+', '').replace('#', '').isalnum():
                        skills.add(skill)
        
        # Method 2: Common technical skills
        tech_skills = [
            'Python', 'JavaScript', 'Java', 'C++', 'C#', 'TypeScript', 'Go', 'Rust',
            'React', 'Angular', 'Vue.js', 'Node.js', 'Django', 'Flask', 'Spring',
            'MongoDB', 'PostgreSQL', 'MySQL', 'Redis', 'AWS', 'Azure', 'Docker',
            'Kubernetes', 'Git', 'Linux', 'REST API', 'GraphQL', 'Machine Learning',
            'TensorFlow', 'PyTorch', 'SQL', 'HTML', 'CSS', 'Bootstrap', 'Tailwind'
        ]
        
        text_lower = text.lower()
        for skill in tech_skills:
            if skill.lower() in text_lower:
                skills.add(skill)
        
        return list(skills)[:20]  # Limit to 20 skills
    
    def _extract_experience(self, text: str) -> List[WorkExperience]:
        """Extract work experience"""
        # Find experience section
        exp_pattern = r'(?:Work\s+|Professional\s+)?Experience[:\s]*\n?(.*?)(?=\n\s*(?:Education|Skills|Projects)\s*:|\Z)'
        match = re.search(exp_pattern, text, re.IGNORECASE | re.MULTILINE | re.DOTALL)
        
        if not match:
            return []
        
        experience_text = match.group(1).strip()
        
        # Split into job blocks
        job_blocks = re.split(r'\n\s*\n', experience_text)
        
        experiences = []
        for block in job_blocks:
            if len(block.strip()) < 20:
                continue
            
            job = self._parse_job_block(block.strip())
            if job:
                experiences.append(job)
        
        return experiences[:5]  # Limit to 5 experiences
    
    def _parse_job_block(self, block: str) -> Optional[WorkExperience]:
        """Parse a single job experience block"""
        lines = [line.strip() for line in block.split('\n') if line.strip()]
        if len(lines) < 2:
            return None
        
        company = ""
        title = ""
        dates = ""
        location = ""
        description = []
        
        # Parse first few lines for company, title, dates
        for i, line in enumerate(lines[:4]):
            # Company indicators
            if any(indicator in line.lower() for indicator in ['inc', 'llc', 'corp', 'company', 'technologies']) and not company:
                company = line
            # Title indicators
            elif any(indicator in line.lower() for indicator in ['engineer', 'developer', 'manager', 'analyst', 'director']) and not title:
                title = line
            # Date patterns
            elif re.search(r'\d{4}|present|current', line, re.IGNORECASE) and not dates:
                dates = line
            # Location patterns
            elif re.search(r'[A-Z][a-z]+,\s*[A-Z]{2}', line) and not location:
                location = line
            else:
                description.append(line)
        
        # Add remaining lines to description
        for line in lines[4:]:
            description.append(line)
        
        # Parse dates
        start_date, end_date = self._parse_dates(dates)
        
        if company or title:
            return WorkExperience(
                company=company or "Unknown Company",
                title=title or "Unknown Position",
                location=location or None,
                start_date=start_date,
                end_date=end_date,
                description='\n'.join(description[:3]) if description else None  # Limit description
            )
        
        return None
    
    def _parse_dates(self, date_text: str) -> tuple[Optional[str], Optional[str]]:
        """Parse start and end dates from text"""
        if not date_text:
            return None, None
        
        # Common date patterns
        patterns = [
            r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4})\s*[-–—to]\s*((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}|present|current)',
            r'(\d{4})\s*[-–—to]\s*(\d{4}|present|current)',
            r'(\d{1,2}/\d{4})\s*[-–—to]\s*(\d{1,2}/\d{4}|present|current)'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, date_text, re.IGNORECASE)
            if match:
                start_date = match.group(1).strip()
                end_date = match.group(2).strip()
                if end_date.lower() in ['present', 'current']:
                    end_date = "Present"
                return start_date, end_date
        
        return None, None
    
    def _extract_education(self, text: str) -> str:
        """Extract education section"""
        patterns = [
            r'Education[:\s]*\n?(.*?)(?=\n\s*(?:Experience|Skills|Projects)\s*:|\Z)',
            r'Academic\s+Background[:\s]*\n?(.*?)(?=\n\s*(?:Experience|Skills|Projects)\s*:|\Z)'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE | re.DOTALL)
            if match:
                education = match.group(1).strip()
                education = re.sub(r'\s+', ' ', education)
                if len(education) > 10:
                    return education[:300]  # Limit length
        
        return ""
    
    def _extract_location(self, text: str) -> str:
        """Extract location using spaCy NER and patterns"""
        # Method 1: spaCy NER
        if self.nlp:
            try:
                doc = self.nlp(text[:1000])
                for ent in doc.ents:
                    if ent.label_ in ["GPE", "LOC"]:
                        location = ent.text.strip()
                        if self._is_valid_location(location):
                            return location
            except Exception as e:
                logger.debug(f"spaCy location extraction failed: {e}")
        
        # Method 2: Pattern-based
        location_patterns = [
            r'(?:Location|Address|Based in)[:\s]*([^\n]+)',
            r'([A-Z][a-z]+,\s*[A-Z]{2})',  # City, ST
            r'([A-Z][a-z]+,\s*[A-Z][a-z]+)'  # City, Country
        ]
        
        for pattern in location_patterns:
            match = re.search(pattern, text)
            if match:
                location = match.group(1).strip()
                if self._is_valid_location(location):
                    return location
        
        return ""
    
    def _is_valid_location(self, location: str) -> bool:
        """Validate if text is likely a location"""
        if not location or len(location) < 3 or len(location) > 50:
            return False
        
        # Should not contain email or URL indicators
        if any(char in location for char in ['@', 'http', 'www']):
            return False
        
        return True
    
    def _extract_linkedin_url(self, text: str) -> Optional[str]:
        """Extract LinkedIn URL"""
        patterns = [
            r'https?://(?:www\.)?linkedin\.com/in/[\w-]+',
            r'linkedin\.com/in/[\w-]+'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                url = match.group(0)
                if not url.startswith('http'):
                    url = 'https://' + url
                return url
        
        return None
    
    def _extract_github_url(self, text: str) -> Optional[str]:
        """Extract GitHub URL"""
        patterns = [
            r'https?://(?:www\.)?github\.com/[\w-]+',
            r'github\.com/[\w-]+'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                url = match.group(0)
                if not url.startswith('http'):
                    url = 'https://' + url
                return url
        
        return None
    
    def _extract_years_experience(self, text: str) -> Optional[int]:
        """Extract years of experience"""
        patterns = [
            r'(\d+)\+?\s*years?\s*(?:of\s*)?experience',
            r'experience[:\s]*(\d+)\+?\s*years?',
            r'(\d+)\+?\s*years?\s*in\s*(?:software|development|engineering)'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                years = int(match.group(1))
                if 0 <= years <= 50:  # Reasonable range
                    return years
        
        return None
